import os
import uuid
import re
import PyPDF2
import docx
from typing import Optional, Dict, Any
from pathlib import Path
import tempfile
import logging
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import SentenceTransformerEmbeddings

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

EMBED = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")
DATA_DIR = Path("data")
DATA_DIR.mkdir(exist_ok=True)

class DocumentProcessor:
    """Handles document processing for various file formats"""

    def __init__(self):
        self.supported_formats = ['.pdf', '.txt', '.docx']

    def process_document(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Process uploaded document and extract text content
        """
        try:
            file_extension = Path(filename).suffix.lower()

            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")

            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
                tmp_file.write(file_content)
                tmp_file_path = tmp_file.name

            try:
                # Extract text based on file type
                if file_extension == '.pdf':
                    text = self._extract_from_pdf(tmp_file_path)
                elif file_extension == '.txt':
                    text = self._extract_from_txt(tmp_file_path)
                elif file_extension == '.docx':
                    text = self._extract_from_docx(tmp_file_path)
                else:
                    raise ValueError(f"Unsupported file format: {file_extension}")

                # Clean and validate text
                text = self._clean_text(text)

                if not text.strip():
                    raise ValueError("No text content found in the document")

                return {
                    "text": text,
                    "filename": filename,
                    "file_type": file_extension,
                    "word_count": len(text.split()),
                    "char_count": len(text),
                    "status": "success"
                }

            finally:
                # Clean up temporary file
                if os.path.exists(tmp_file_path):
                    os.unlink(tmp_file_path)

        except Exception as e:
            logger.error(f"Error processing document {filename}: {str(e)}")
            return {
                "text": "",
                "filename": filename,
                "error": str(e),
                "status": "error"
            }

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.error(f"Error extracting from PDF: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
        return text

    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
            except Exception as e:
                logger.error(f"Error reading TXT file: {str(e)}")
                raise ValueError(f"Failed to read TXT file: {str(e)}")
        except Exception as e:
            logger.error(f"Error extracting from TXT: {str(e)}")
            raise ValueError(f"Failed to extract text from TXT: {str(e)}")
        return text

    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            logger.error(f"Error extracting from DOCX: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
        return text

    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        text = ' '.join(text.split())
        text = text.replace('\x00', '')
        text = text.replace('\r', '\n')
        while '\n\n\n' in text:
            text = text.replace('\n\n\n', '\n\n')
        return text.strip()

    def get_text_chunks(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> list:
        """
        Split text into chunks for processing
        """
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            if end < len(text):
                sentence_end = text.rfind('.', start, end)
                if sentence_end > start + chunk_size - 100:
                    end = sentence_end + 1
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            start = end - overlap
            if start >= len(text):
                break
        return chunks

def load_text(file_path: Path) -> str:
    # Add support for PDF and TXT
    if file_path.suffix == ".pdf":
        reader = PyPDF2.PdfReader(str(file_path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    elif file_path.suffix == ".txt":
        return file_path.read_text()
    elif file_path.suffix == ".docx":
        doc = docx.Document(str(file_path))
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)
    else:
        raise ValueError("Unsupported file type")

def build_vector_index(text: str) -> FAISS:
    # Increase chunk size for research papers
    splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=400)
    docs = splitter.create_documents([text])
    return FAISS.from_documents(docs, EMBED)

def save_doc_and_index(file_bytes: bytes, filename: str):
    file_path = DATA_DIR / filename
    file_path.write_bytes(file_bytes)
    text = load_text(file_path)
    index = build_vector_index(text)
    idx_path = DATA_DIR / f"{file_path.stem}.faiss"
    index.save_local(idx_path)
    return file_path.stem, text

def summarize(text: str) -> str:
    # Simple summary: first 3 sentences
    return " ".join(text.split(".")[:3]) + "..."